import io
from docx import Document
from app.services.ingestion.text_extraction import extract_text


def test_extract_text_passthrough_strips():
    assert extract_text("text", text_value="  Hello world  ") == "Hello world"


def test_extract_text_empty_raises():
    import pytest
    with pytest.raises(ValueError):
        extract_text("text", text_value="   ")


def test_extract_docx():
    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("First line of the sermon.")
    doc.add_paragraph("Second line about grace.")
    doc.save(buf)
    out = extract_text("docx", file_bytes=buf.getvalue())
    assert "First line of the sermon." in out
    assert "grace" in out
