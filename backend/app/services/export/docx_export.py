"""
Export service — generates .docx manuscripts from chapter HTML content.
"""
import io
import re
from typing import List, Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from html.parser import HTMLParser


def html_to_plain(html: str) -> str:
    """Strip HTML tags to plain text for docx export."""
    text = re.sub(r'<br\s*/?>', '\n', html)
    text = re.sub(r'<p[^>]*>', '\n', text)
    text = re.sub(r'</p>', '', text)
    text = re.sub(r'<[^>]+>', '', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


class DocxHtmlParser(HTMLParser):
    def __init__(self, doc: Document):
        super().__init__()
        self.doc = doc
        self.current_paragraph = None
        self.is_bold = False
        self.is_italic = False
        self.is_underline = False
        self.in_blockquote = False
        self.list_type = None  # 'ul' or 'ol'

    def handle_starttag(self, tag, attrs):
        if tag in ('p', 'h1', 'h2', 'h3', 'blockquote', 'li'):
            if tag == 'h1':
                self.current_paragraph = self.doc.add_paragraph()
                run = self.current_paragraph.add_run()
                run.bold = True
                run.font.size = Pt(16)
                self.current_paragraph.paragraph_format.space_before = Pt(12)
                self.current_paragraph.paragraph_format.space_after = Pt(6)
            elif tag == 'h2':
                self.current_paragraph = self.doc.add_paragraph()
                run = self.current_paragraph.add_run()
                run.bold = True
                run.font.size = Pt(14)
                self.current_paragraph.paragraph_format.space_before = Pt(12)
                self.current_paragraph.paragraph_format.space_after = Pt(4)
            elif tag == 'h3':
                self.current_paragraph = self.doc.add_paragraph()
                run = self.current_paragraph.add_run()
                run.bold = True
                run.font.size = Pt(12)
                self.current_paragraph.paragraph_format.space_before = Pt(6)
                self.current_paragraph.paragraph_format.space_after = Pt(2)
            elif tag == 'blockquote':
                self.current_paragraph = self.doc.add_paragraph()
                self.current_paragraph.paragraph_format.left_indent = Inches(0.5)
                self.current_paragraph.paragraph_format.space_after = Pt(6)
                self.in_blockquote = True
            elif tag == 'li':
                style = 'List Bullet' if self.list_type == 'ul' else 'List Number'
                self.current_paragraph = self.doc.add_paragraph(style=style)
                self.current_paragraph.paragraph_format.space_after = Pt(3)
            else:
                self.current_paragraph = self.doc.add_paragraph()
                self.current_paragraph.paragraph_format.first_line_indent = Inches(0.5)
                self.current_paragraph.paragraph_format.space_after = Pt(6)
        elif tag in ('strong', 'b'):
            self.is_bold = True
        elif tag in ('em', 'i'):
            self.is_italic = True
        elif tag in ('u',):
            self.is_underline = True
        elif tag in ('ul', 'ol'):
            self.list_type = tag

    def handle_endtag(self, tag):
        if tag in ('p', 'h1', 'h2', 'h3', 'blockquote', 'li'):
            self.current_paragraph = None
            if tag == 'blockquote':
                self.in_blockquote = False
        elif tag in ('strong', 'b'):
            self.is_bold = False
        elif tag in ('em', 'i'):
            self.is_italic = False
        elif tag in ('u',):
            self.is_underline = False
        elif tag in ('ul', 'ol'):
            self.list_type = None

    def handle_data(self, data):
        if not self.current_paragraph:
            self.current_paragraph = self.doc.add_paragraph()
            self.current_paragraph.paragraph_format.first_line_indent = Inches(0.5)
            self.current_paragraph.paragraph_format.space_after = Pt(6)

        run = self.current_paragraph.add_run(data)
        run.font.name = 'Times New Roman'

        # Apply formatting
        if self.is_bold:
            run.bold = True
        if self.is_underline:
            run.underline = True

        # Check if first run in paragraph is scripture and paragraph style is normal (not heading/list/blockquote)
        is_first_run = len(self.current_paragraph.runs) == 1
        is_scripture = False
        if is_first_run and not self.in_blockquote and not self.list_type:
            if re.match(r'^[A-Z][a-zA-Z\s]+ \d+:\d+', data.strip()):
                is_scripture = True

        if self.is_italic or self.in_blockquote or is_scripture:
            run.italic = True

        # Size formatting (unless it's set by starttag heading initialization)
        is_heading = len(self.current_paragraph.runs) > 1 and self.current_paragraph.runs[0].bold and self.current_paragraph.runs[0].font.size is not None
        if is_heading:
            run.font.size = self.current_paragraph.runs[0].font.size
            run.bold = True
        else:
            run.font.size = Pt(11 if (self.in_blockquote or is_scripture) else 12)


def create_manuscript_docx(
    title: str,
    author_name: str,
    chapters: List[dict],  # [{number, title, content}]
) -> bytes:
    """
    Generate a publisher-ready .docx manuscript.
    """
    doc = Document()

    # Page setup: 1-inch margins, 12pt Times New Roman
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)

    # Title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title.upper())
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.name = 'Times New Roman'

    doc.add_paragraph()
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_para.add_run(f"by {author_name}")
    author_run.font.size = Pt(14)
    author_run.font.name = 'Times New Roman'

    doc.add_page_break()

    # Chapters
    for chapter in chapters:
        # Chapter heading
        ch_para = doc.add_paragraph()
        ch_run = ch_para.add_run(f"Chapter {chapter['number']}")
        ch_run.bold = True
        ch_run.font.size = Pt(11)
        ch_run.font.name = 'Times New Roman'

        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t_run = title_p.add_run(chapter['title'].upper())
        t_run.bold = True
        t_run.font.size = Pt(14)
        t_run.font.name = 'Times New Roman'

        doc.add_paragraph()

        # Parse and add chapter content HTML
        content_html = chapter.get('content') or ''
        
        # If content has no HTML tags, wrap it in paragraphs
        if not content_html.strip().startswith('<'):
            content_html = '\n\n'.join(f'<p>{p}</p>' for p in content_html.split('\n\n'))
            
        parser = DocxHtmlParser(doc)
        parser.feed(content_html)

        doc.add_page_break()

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
